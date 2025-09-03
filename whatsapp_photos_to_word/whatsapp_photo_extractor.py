#!/usr/bin/env python3
"""
WhatsApp Chat Photo Extractor
Extracts photos from a WhatsApp chat export zip file and creates a Word document
with photos, timestamps, senders, and associated messages as captions.
"""

import zipfile
import re
import os
import tempfile
import shutil
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse

def parse_message_date(date_str):
    """Parse WhatsApp date string to datetime object."""
    try:
        # WhatsApp format: dd/mm/yy or dd/mm/yyyy
        parts = date_str.split('/')
        if len(parts) == 3:
            day, month, year = parts
            # Convert 2-digit year to 4-digit
            if len(year) == 2:
                year = '20' + year if int(year) < 50 else '19' + year
            
            return datetime(int(year), int(month), int(day))
    except:
        pass
    return None

def parse_chat_messages(chat_content):
    """Parse WhatsApp chat content and extract messages with metadata."""
    # Decode content with UTF-8 encoding
    if isinstance(chat_content, bytes):
        chat_content = chat_content.decode('utf-8', errors='replace')
    
    messages = []
    lines = chat_content.split('\n')
    
    # Pattern to match WhatsApp message format: "date, time - sender: message"
    message_pattern = r'^(\d{1,2}/\d{1,2}/\d{2,4}), (\d{1,2}:\d{2}(?:\u202f)?(?:am|pm)?) - ([^:]+): (.+)$'
    
    current_message = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this is a new message
        match = re.match(message_pattern, line)
        if match:
            date_str, time_str, sender, content = match.groups()
            
            # Clean up the time string (remove narrow no-break space)
            time_str = time_str.replace('\u202f', ' ')
            
            current_message = {
                'date': date_str,
                'time': time_str,
                'sender': sender.strip(),
                'content': content.strip(),
                'full_content': content.strip(),
                'datetime': parse_message_date(date_str)
            }
            messages.append(current_message)
        else:
            # This is a continuation of the previous message
            if current_message:
                current_message['full_content'] += '\n' + line
    
    return messages

def find_photo_messages(messages):
    """Find messages that contain photo attachments."""
    photo_messages = []
    
    for message in messages:
        # Look for image file references
        if '(file attached)' in message['content'] or '(file attached)' in message['full_content']:
            # Extract image filename
            img_pattern = r'(IMG-\d{8}-WA\d{4}\.jpg)'
            match = re.search(img_pattern, message['content'])
            if match:
                image_filename = match.group(1)
                message['image_filename'] = image_filename
                photo_messages.append(message)
    
    return photo_messages

def filter_messages_by_date(messages, start_date=None, end_date=None):
    """Filter messages by date range."""
    if not start_date and not end_date:
        return messages
    
    filtered_messages = []
    for message in messages:
        msg_date = message.get('datetime')
        if not msg_date:
            continue  # Skip messages with unparseable dates
            
        # Check if message falls within date range
        include_message = True
        
        if start_date and msg_date < start_date:
            include_message = False
        
        if end_date and msg_date > end_date:
            include_message = False
            
        if include_message:
            filtered_messages.append(message)
    
    return filtered_messages

def get_person_initials(name):
    """Extract initials from person's name."""
    # Split by spaces and take first letter of each word
    words = name.strip().split()
    initials = ''.join([word[0].upper() for word in words if word])
    return initials[:3]  # Limit to 3 characters

def extract_descriptive_words(text, max_chars=15):
    """Extract the most relevant descriptive words from text."""
    # Remove file attachment references
    text = re.sub(r'IMG-\d{8}-WA\d{4}\.jpg \(file attached\)', '', text)
    text = text.strip()
    
    if not text:
        return 'NoText'
    
    # Common words to filter out (keep it concise for filenames)
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 
        'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be', 'been', 'have', 'has', 'had',
        'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must',
        'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them',
        'my', 'your', 'his', 'her', 'its', 'our', 'their', 'this', 'that', 'these', 'those'
    }
    
    # Clean and split into words
    words = re.findall(r'\b\w+\b', text.lower())
    
    # Filter out stop words and very short words, prioritize longer meaningful words
    descriptive_words = []
    for word in words:
        if len(word) >= 2 and word not in stop_words:
            descriptive_words.append(word)
    
    # If no descriptive words found, use first few words
    if not descriptive_words:
        words = re.findall(r'\b\w+\b', text)
        descriptive_words = words[:3] if words else ['NoText']
    
    # Join words and truncate to max_chars
    result = ' '.join(descriptive_words)
    if len(result) > max_chars:
        # Try to cut at word boundary if possible
        truncated = result[:max_chars]
        last_space = truncated.rfind(' ')
        if last_space > max_chars * 0.7:  # If we can cut at a word boundary reasonably close
            result = truncated[:last_space]
        else:
            result = truncated
    else:
        result = result
    
    return result if result else 'NoText'

def sanitize_text_for_filename(text):
    """Clean text to be safe for filename."""
    descriptive_text = extract_descriptive_words(text, 15)
    
    # Remove or replace invalid filename characters
    descriptive_text = re.sub(r'[<>:"/\\|?*]', '', descriptive_text)
    descriptive_text = re.sub(r'[\r\n\t]', ' ', descriptive_text)
    descriptive_text = re.sub(r'\s+', ' ', descriptive_text)  # Multiple spaces to single space
    
    return descriptive_text.strip() if descriptive_text.strip() else 'NoText'

def generate_new_filename(message):
    """Generate new filename based on date, initials, and message text."""
    # Parse date (format: dd/mm/yy)
    try:
        date_parts = message['date'].split('/')
        if len(date_parts) == 3:
            day, month, year = date_parts
            # Convert 2-digit year to full year
            if len(year) == 2:
                year = '20' + year if int(year) < 50 else '19' + year
            
            # Format as YYMMDD
            date_formatted = f"{year[-2:]}{month.zfill(2)}{day.zfill(2)}"
        else:
            date_formatted = "000000"
    except:
        date_formatted = "000000"
    
    # Get initials
    initials = get_person_initials(message['sender'])
    
    # Get first 10 characters of message text
    text_snippet = sanitize_text_for_filename(message['full_content'])
    
    # Original extension
    original_filename = message['image_filename']
    extension = os.path.splitext(original_filename)[1]
    
    # Create new filename
    new_filename = f"{date_formatted}_{initials}_{text_snippet}{extension}"
    
    return new_filename

def extract_photos(photo_messages, zip_file_path, extract_dir):
    """Extract photos from zip and rename them according to the new format."""
    if not os.path.exists(extract_dir):
        os.makedirs(extract_dir)
    
    extracted_files = []
    
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        for message in photo_messages:
            if 'image_filename' not in message:
                continue
                
            original_filename = message['image_filename']
            new_filename = generate_new_filename(message)
            
            try:
                # Extract to temporary location
                zip_ref.extract(original_filename, extract_dir)
                temp_path = os.path.join(extract_dir, original_filename)
                
                # Rename to new format
                new_path = os.path.join(extract_dir, new_filename)
                
                # Handle duplicate filenames
                counter = 1
                base_new_path = new_path
                while os.path.exists(new_path):
                    name, ext = os.path.splitext(base_new_path)
                    new_path = f"{name}_{counter}{ext}"
                    counter += 1
                
                shutil.move(temp_path, new_path)
                extracted_files.append({
                    'original': original_filename,
                    'new': os.path.basename(new_path),
                    'path': new_path,
                    'message': message
                })
                
            except Exception as e:
                print(f"Error extracting {original_filename}: {str(e)}")
                continue
    
    return extracted_files

def create_word_document(photo_messages, zip_file_path, output_path, start_date=None, end_date=None):
    """Create a Word document with photos and captions."""
    doc = Document()
    
    # Add title with date range if filtered
    title_text = 'WhatsApp Photo Report'
    if start_date or end_date:
        if start_date and end_date:
            title_text += f' ({start_date.strftime("%d/%m/%Y")} - {end_date.strftime("%d/%m/%Y")})'
        elif start_date:
            title_text += f' (from {start_date.strftime("%d/%m/%Y")})'
        elif end_date:
            title_text += f' (until {end_date.strftime("%d/%m/%Y")})'
    
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation timestamp and period info
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    
    if start_date or end_date:
        period_text = "Period: "
        if start_date and end_date:
            period_text += f"{start_date.strftime('%B %d, %Y')} to {end_date.strftime('%B %d, %Y')}"
        elif start_date:
            period_text += f"From {start_date.strftime('%B %d, %Y')} onwards"
        elif end_date:
            period_text += f"Until {end_date.strftime('%B %d, %Y')}"
        doc.add_paragraph(period_text)
    
    doc.add_paragraph()
    
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        # Create temporary directory for extracted images
        with tempfile.TemporaryDirectory() as temp_dir:
            photo_count = 0
            
            for message in photo_messages:
                if 'image_filename' not in message:
                    continue
                
                image_filename = message['image_filename']
                
                try:
                    # Extract the image to temporary directory
                    zip_ref.extract(image_filename, temp_dir)
                    image_path = os.path.join(temp_dir, image_filename)
                    
                    # Add photo to document
                    photo_count += 1
                    
                    # Add section heading
                    doc.add_heading(f'Photo {photo_count}', level=2)
                    
                    # Add the image (resize to fit page)
                    try:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(image_path, width=Inches(3))
                    except Exception as e:
                        doc.add_paragraph(f"Error loading image {image_filename}: {str(e)}")
                        continue
                    
                    # Add caption with message details
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Format the caption
                    caption_text = f"From: {message['sender']}\n"
                    caption_text += f"Date: {message['date']} at {message['time']}\n"
                    
                    # Add associated message text (excluding the file attachment line)
                    message_text = message['full_content']
                    # Remove the "IMG-xxx.jpg (file attached)" line
                    message_text = re.sub(r'IMG-\d{8}-WA\d{4}\.jpg \(file attached\)', '', message_text)
                    message_text = message_text.strip()
                    
                    if message_text:
                        caption_text += f"Message: {message_text}"
                    
                    caption_para.add_run(caption_text).italic = True
                    
                    # Add some spacing
                    doc.add_paragraph()
                    
                except KeyError:
                    print(f"Warning: Image {image_filename} not found in zip file")
                    continue
                except Exception as e:
                    print(f"Error processing image {image_filename}: {str(e)}")
                    continue
    
    # Save the document
    doc.save(output_path)
    return photo_count

def interactive_mode():
    """Interactive mode for user-friendly input."""
    print("=== WhatsApp Photo Extractor ===")
    print()
    
    # Get zip file
    zip_file = input("Enter WhatsApp zip file name (or press Enter for 'WhatsApp Chat.zip'): ").strip()
    if not zip_file:
        zip_file = "WhatsApp Chat.zip"
    
    # Get date filtering option
    print("\nDate filtering options:")
    print("1. All photos (no date filter)")
    print("2. Last month only")  
    print("3. Custom date range")
    
    choice = input("Choose option (1-3): ").strip()
    
    start_date_str = None
    end_date_str = None
    last_month = False
    
    if choice == "2":
        last_month = True
    elif choice == "3":
        start_date_str = input("Start date (DD/MM/YYYY, or press Enter to skip): ").strip()
        if not start_date_str:
            start_date_str = None
        end_date_str = input("End date (DD/MM/YYYY, or press Enter to skip): ").strip()
        if not end_date_str:
            end_date_str = None
    
    # Get output options
    output_doc = input("\nOutput Word document name (or press Enter for 'whatsapp_photos.docx'): ").strip()
    if not output_doc:
        output_doc = "whatsapp_photos.docx"
    
    extract_dir = input("Photo extraction folder (or press Enter for 'extracted_photos'): ").strip()
    if not extract_dir:
        extract_dir = "extracted_photos"
    
    extract_only = input("Extract photos only, no Word document? (y/N): ").strip().lower()
    extract_only = extract_only.startswith('y')
    
    return {
        'zip_file': zip_file,
        'output': output_doc,
        'extract_dir': extract_dir,
        'extract_only': extract_only,
        'start_date': start_date_str,
        'end_date': end_date_str,
        'last_month': last_month
    }

def main():
    parser = argparse.ArgumentParser(description='Extract photos from WhatsApp chat export and create Word document')
    parser.add_argument('zip_file', nargs='?', default=None, 
                       help='Path to WhatsApp chat export zip file (default: WhatsApp Chat.zip)')
    parser.add_argument('-o', '--output', default='whatsapp_photos.docx', 
                       help='Output Word document path (default: whatsapp_photos.docx)')
    parser.add_argument('-e', '--extract-dir', default='extracted_photos', 
                       help='Directory to extract renamed photos (default: extracted_photos)')
    parser.add_argument('--extract-only', action='store_true',
                       help='Only extract photos, do not create Word document')
    parser.add_argument('--start-date', type=str,
                       help='Start date for filtering (format: DD/MM/YYYY or DD/MM/YY)')
    parser.add_argument('--end-date', type=str,
                       help='End date for filtering (format: DD/MM/YYYY or DD/MM/YY)')
    parser.add_argument('--last-month', action='store_true',
                       help='Filter to show only photos from the last month')
    parser.add_argument('--interactive', action='store_true',
                       help='Run in interactive mode with prompts')
    
    args = parser.parse_args()
    
    # If no zip file specified and no other arguments, run interactive mode
    if args.zip_file is None and not any([args.start_date, args.end_date, args.last_month, args.interactive]):
        args.interactive = True
    
    # Run interactive mode if requested or needed
    if args.interactive:
        interactive_args = interactive_mode()
        # Override args with interactive input
        args.zip_file = interactive_args['zip_file']
        args.output = interactive_args['output']
        args.extract_dir = interactive_args['extract_dir']
        args.extract_only = interactive_args['extract_only']
        args.start_date = interactive_args['start_date']
        args.end_date = interactive_args['end_date']
        args.last_month = interactive_args['last_month']
    
    # Set default zip file if still not specified
    if args.zip_file is None:
        args.zip_file = "WhatsApp Chat.zip"
    
    # Parse date filters
    start_date = None
    end_date = None
    
    if args.last_month:
        # Calculate date range for last month
        today = datetime.now()
        end_date = today
        start_date = today - timedelta(days=30)
        print(f"Filtering for last month: {start_date.strftime('%d/%m/%Y')} to {end_date.strftime('%d/%m/%Y')}")
    else:
        if args.start_date:
            start_date = parse_message_date(args.start_date)
            if not start_date:
                print(f"Error: Invalid start date format '{args.start_date}'. Use DD/MM/YYYY or DD/MM/YY")
                return 1
            print(f"Start date filter: {start_date.strftime('%d/%m/%Y')}")
        
        if args.end_date:
            end_date = parse_message_date(args.end_date)
            if not end_date:
                print(f"Error: Invalid end date format '{args.end_date}'. Use DD/MM/YYYY or DD/MM/YY")
                return 1
            print(f"End date filter: {end_date.strftime('%d/%m/%Y')}")
    
    if not os.path.exists(args.zip_file):
        print(f"Error: Zip file '{args.zip_file}' not found")
        return 1
    
    print(f"Processing WhatsApp chat export: {args.zip_file}")
    
    try:
        # Open zip file and read chat content
        with zipfile.ZipFile(args.zip_file, 'r') as zip_ref:
            # Find the chat text file (usually ends with .txt)
            txt_files = [f for f in zip_ref.namelist() if f.endswith('.txt')]
            if not txt_files:
                print("Error: No chat text file found in zip")
                return 1
            
            chat_file = txt_files[0]  # Take the first .txt file
            print(f"Reading chat from: {chat_file}")
            
            chat_content = zip_ref.read(chat_file)
            
        # Parse messages
        print("Parsing chat messages...")
        messages = parse_chat_messages(chat_content)
        print(f"Found {len(messages)} total messages")
        
        # Apply date filters if specified
        if start_date or end_date:
            filtered_messages = filter_messages_by_date(messages, start_date, end_date)
            print(f"After date filtering: {len(filtered_messages)} messages")
            messages = filtered_messages
        
        # Find photo messages
        photo_messages = find_photo_messages(messages)
        print(f"Found {len(photo_messages)} messages with photos")
        
        if not photo_messages:
            print("No photo messages found in chat")
            return 0
        
        # Extract photos with new names
        print(f"Extracting photos to: {args.extract_dir}")
        extracted_files = extract_photos(photo_messages, args.zip_file, args.extract_dir)
        print(f"Successfully extracted {len(extracted_files)} photos")
        
        # Show some examples of the new filenames
        if extracted_files:
            print("\nExample renamed files:")
            for i, file_info in enumerate(extracted_files[:5]):  # Show first 5
                print(f"  {file_info['original']} -> {file_info['new']}")
            if len(extracted_files) > 5:
                print(f"  ... and {len(extracted_files) - 5} more files")
        
        # Create Word document unless extract-only mode
        if not args.extract_only:
            print(f"\nCreating Word document: {args.output}")
            photo_count = create_word_document(photo_messages, args.zip_file, args.output, start_date, end_date)
            print(f"Successfully created document with {photo_count} photos")
            print(f"Output saved to: {args.output}")
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return 1
    
    return 0

if __name__ == '__main__':
    import sys
    sys.exit(main())